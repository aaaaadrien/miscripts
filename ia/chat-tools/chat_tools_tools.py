"""
chat_tools_tools.py
===================
Module partagé contenant les outils (météo, wiki, change).
Importé par chat-tools-cli.py et chat-tools-web.py.
Important : pas de - mais des _ dans le fichier pour importer le module !

Utilisation :
    from chat_tools_tools import CATALOGUE_OUTILS, outils_actifs, executer_outil
"""

import configparser
import re
from datetime import datetime

import requests
from bs4 import BeautifulSoup
from ddgs import DDGS

# Récupère le bulletin météo complet pour une ville donnée.
def outil_meteo(ville: str) -> str:
    """
      1. Géolocalisation via l'API open-meteo geocoding
      2. Récupération des données météo (température, humidité, vent, état du ciel)
      3. Traduction du code météo WMO en texte lisible
    """
    try:
        geo_url = (
            f"https://geocoding-api.open-meteo.com/v1/search"
            f"?name={ville}&count=1&language=fr&format=json"
        )
        geo = requests.get(geo_url, timeout=5).json()
        if not geo.get("results"):
            return f"Désolé, la ville « {ville} » est introuvable."

        lieu = geo["results"][0]
        lat, lon = lieu["latitude"], lieu["longitude"]

        meteo_url = (
            f"https://api.open-meteo.com/v1/forecast"
            f"?latitude={lat}&longitude={lon}"
            f"&current=temperature_2m,relative_humidity_2m,weather_code,wind_speed_10m"
            f"&timezone=auto"
        )
        res     = requests.get(meteo_url, timeout=5).json()
        donnees = res["current"]

        traduction_temps = {
            0:  "Ciel dégagé ☀️",    1:  "Principalement dégagé 🌤️",
            2:  "Partiellement nuageux ⛅", 3: "Couvert ☁️",
            45: "Brouillard 🌫️",     48: "Brouillard givrant 🌫️❄️",
            51: "Bruine légère 🌦️",  61: "Pluie légère 🌧️",
            63: "Pluie modérée 🌧️",  65: "Pluie forte 🌧️💧",
            71: "Chute de neige légère 🌨️", 95: "Orage ⛈️",
        }
        etat = traduction_temps.get(donnees["weather_code"], "Conditions variables")

        return (
            f"Météo à {lieu['name']} ({lieu.get('country', '')}) : {etat}\n"
            f"  🌡 Température : {donnees['temperature_2m']}°C\n"
            f"  💧 Humidité    : {donnees['relative_humidity_2m']}%\n"
            f"  💨 Vent        : {donnees['wind_speed_10m']} km/h"
        )

    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé pour l'API météo."
    except Exception as e:
        return f"⚠️ Impossible de récupérer la météo ({e})."

# Retourne le résumé Wikipédia (version française) du sujet demandé.
def outil_wiki(sujet: str) -> str:
    """
    Utilise l'API REST de Wikipédia pour obtenir l'extrait de la page.
    """

    # Fix sinon bloqué
    headers = {
        'User-Agent': 'Linuxtricks/1.0'
    }
    
    try:
        url  = f"https://fr.wikipedia.org/api/rest_v1/page/summary/{sujet.replace(' ', '_')}"
        res  = requests.get(url, headers=headers, timeout=5).json()
        texte = res.get("extract")
        if not texte:
            return f"Aucune page Wikipédia trouvée pour « {sujet} »."
        return texte

    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé pour Wikipédia."
    except Exception as e:
        return f"⚠️ Erreur lors de la recherche Wikipédia ({e})."
        
# Retourne la page Wikipédia (version française) du sujet demandé.
def outil_wiki_full(sujet: str) -> str:
    """
    Utilise l'Action API de Wikipédia pour obtenir l'entièreté du contenu d'une page.
    """
    url = "https://fr.wikipedia.org/w/api.php"
    
    # Fix sinon bloqué
    headers = {
        'User-Agent': 'Linuxtricks/1.0'
    }
    
    params = {
        "action": "query",
        "prop": "extracts",
        "explaintext": True,  # Récupère le texte brut sans HTML
        "titles": sujet,
        "format": "json",
        "redirects": 1
    }

    try:
        response = requests.get(url, params=params, headers=headers, timeout=5)
        # Vérifie si la requête a réussi avant de tenter le décodage JSON
        response.raise_for_status()
        
        data = response.json()
        pages = data.get("query", {}).get("pages", {})
        
        # Récupération de la première page trouvée
        page_id = next(iter(pages))
        page_data = pages[page_id]

        if "missing" in page_data or page_id == "-1":
            return f"Aucune page Wikipédia trouvée pour « {sujet} »."

        texte = page_data.get("extract")
        if not texte:
            return f"Le contenu de la page « {sujet} » est vide."
            
        return texte

    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé pour Wikipédia."
    except requests.exceptions.HTTPError as e:
        return f"⚠️ Erreur HTTP (Accès refusé ou page inexistante) : {e}"
    except Exception as e:
        return f"⚠️ Erreur lors de la recherche Wikipédia ({e})."
        

# Convertit une somme d'une devise vers une autre en temps réel.
def outil_argent(montant: float, de_monnaie: str, vers_monnaie: str) -> str:
    """
    Utilise l'API open.er-api.com (gratuite, sans clé).

    Paramètres :
      montant      : valeur numérique à convertir
      de_monnaie   : code ISO 4217 source (ex : EUR)
      vers_monnaie : code ISO 4217 cible  (ex : USD)
    """
    try:
        url = f"https://open.er-api.com/v6/latest/{de_monnaie.upper()}"
        res = requests.get(url, timeout=5).json()

        if res.get("result") != "success":
            return f"⚠️ Devise source « {de_monnaie} » non reconnue."

        taux = res["rates"].get(vers_monnaie.upper())
        if taux is None:
            return f"⚠️ Devise cible « {vers_monnaie} » non reconnue."

        resultat = montant * taux
        return (
            f"💱 {montant} {de_monnaie.upper()} = {resultat:.2f} {vers_monnaie.upper()}\n"
            f"  Taux : 1 {de_monnaie.upper()} = {taux:.4f} {vers_monnaie.upper()}"
        )

    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé pour l'API de change."
    except Exception as e:
        return f"⚠️ Erreur de conversion monétaire ({e})."


# Effectue une recherche web via DuckDuckGo et retourne les N premiers résultats.
def outil_duckduckgo(query: str, num_results: int = 5) -> str:
    """
    Utilise DuckDuckGo (ddgs) pour rechercher sur le web.
    Retourne titres, URLs et extraits pour chaque résultat.
    """
    num_results = min(int(num_results), 10)
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=num_results))
        if not results:
            return "Aucun résultat trouvé."
        lines = [f"**Résultats pour « {query} »**\n"]
        for i, r in enumerate(results, 1):
            lines.append(
                f"{i}. **{r.get('title', 'Sans titre')}**\n"
                f"   URL : {r.get('href', '')}\n"
                f"   {r.get('body', '')}\n"
            )
        return "\n".join(lines)
    except Exception as e:
        return f"⚠️ Erreur lors de la recherche : {e}"


# Télécharge et extrait le texte principal d'une page web.
def outil_recup_page(url: str, max_chars: int = 3000) -> str:
    """
    Récupère une page web via requests + BeautifulSoup.
    Supprime les balises inutiles (nav, scripts, pubs…) et nettoie le texte.
    """
    max_chars = min(int(max_chars), 8000)
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
        )
    }
    try:
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        text = re.sub(r"\n{3,}", "\n\n", text)
        if len(text) > max_chars:
            text = text[:max_chars] + f"\n\n[… contenu tronqué à {max_chars} caractères]"
        return text or "Page vide ou contenu non extractible."
    except requests.exceptions.Timeout:
        return "⚠️ Délai d'attente dépassé lors du chargement de la page."
    except Exception as e:
        return f"⚠️ Erreur lors du chargement de la page : {e}"


# Retourne la date et l'heure actuelles.
def outil_datetime() -> str:
    """
    Retourne la date et l'heure locales formatées en français.
    Aucun paramètre requis.
    """
    return datetime.now().strftime("Date : %A %d %B %Y — Heure : %H:%M:%S")


# Génère un fichier (txt, md, pdf, docx, xlsx) et retourne son contenu en base64.
def outil_generer_fichier(contenu: str, format: str, nom_fichier: str) -> str:
    """
    Génère un fichier téléchargeable dans le format demandé.

    Paramètres :
      contenu      : texte à mettre dans le fichier (markdown accepté)
      format       : extension cible — txt | md | pdf | docx | xlsx
      nom_fichier  : nom du fichier sans extension (ex : rapport_meteo)
    Retourne un JSON avec {b64, mime, nom, format, __fichier_genere__: true}.
    """
    import io, base64 as _b64, json as _json

    fmt = format.lower().lstrip(".")

    # Nom propre
    if "." not in nom_fichier:
        nom_fichier = f"{nom_fichier}.{fmt}"

    try:
        # TXT / MD
        if fmt in ("txt", "md"):
            data = contenu.encode("utf-8")
            mime = "text/plain" if fmt == "txt" else "text/markdown"

        # PDF (fpdf2)
        elif fmt == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.set_margins(15, 15, 15)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            for line in contenu.split("\n"):
                stripped = line.rstrip()
                if stripped == "":
                    pdf.ln(4)
                elif stripped.startswith("### "):
                    pdf.set_font("Helvetica", "B", 12)
                    pdf.multi_cell(0, 7, stripped[4:])
                    pdf.set_font("Helvetica", size=11)
                elif stripped.startswith("## "):
                    pdf.set_font("Helvetica", "B", 14)
                    pdf.multi_cell(0, 8, stripped[3:])
                    pdf.set_font("Helvetica", size=11)
                elif stripped.startswith("# "):
                    pdf.set_font("Helvetica", "B", 16)
                    pdf.multi_cell(0, 10, stripped[2:])
                    pdf.set_font("Helvetica", size=11)
                elif stripped.startswith(("- ", "* ", "• ")):
                    pdf.set_font("Helvetica", size=11)
                    pdf.multi_cell(0, 6, "  • " + stripped[2:])
                else:
                    pdf.set_font("Helvetica", size=11)
                    pdf.multi_cell(0, 6, stripped)

            data = bytes(pdf.output())
            mime = "application/pdf"

        # DOCX (python-docx)
        elif fmt == "docx":
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            for line in contenu.split("\n"):
                stripped = line.rstrip()
                if stripped.startswith("### "):
                    doc.add_heading(stripped[4:], level=3)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=2)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], level=1)
                elif stripped.startswith(("- ", "* ", "• ")):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                elif stripped == "":
                    doc.add_paragraph("")
                else:
                    doc.add_paragraph(stripped)

            buf = io.BytesIO()
            doc.save(buf)
            data = buf.getvalue()
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        # XLSX (openpyxl)
        elif fmt == "xlsx":
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment

            wb = openpyxl.Workbook()
            ws = wb.active
            lines = [l for l in contenu.split("\n")]

            # Détection tableau markdown (lignes avec |)
            table_lines = [l for l in lines if "|" in l and l.strip().startswith("|")]

            if table_lines:
                row_idx = 1
                header_done = False
                for line in lines:
                    if not line.strip():
                        continue
                    if "|" not in line:
                        continue
                    # Ligne séparatrice |---|---|
                    if all(c in "-|: " for c in line):
                        header_done = True
                        continue
                    cells = [c.strip() for c in line.strip().strip("|").split("|")]
                    for col_idx, cell in enumerate(cells, 1):
                        c = ws.cell(row=row_idx, column=col_idx, value=cell)
                        if not header_done:  # ligne d'en-tête
                            c.font = Font(bold=True)
                            c.fill = PatternFill("solid", fgColor="4472C4")
                            c.font = Font(bold=True, color="FFFFFF")
                        c.alignment = Alignment(wrap_text=True)
                    row_idx += 1
                # Auto-largeur approximative
                for col in ws.columns:
                    max_len = max((len(str(c.value or "")) for c in col), default=10)
                    ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)
            else:
                # Texte brut : une ligne = une cellule A?
                for row_idx, line in enumerate(lines, 1):
                    ws.cell(row=row_idx, column=1, value=line)

            buf = io.BytesIO()
            wb.save(buf)
            data = buf.getvalue()
            mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

        else:
            return f"⚠️ Format « {format} » non supporté. Formats disponibles : txt, md, pdf, docx, xlsx"

        b64 = _b64.b64encode(data).decode("utf-8")
        return _json.dumps({
            "__fichier_genere__": True,
            "b64":    b64,
            "mime":   mime,
            "nom":    nom_fichier,
            "format": fmt,
        })

    except ImportError as e:
        return f"⚠️ Bibliothèque manquante pour générer le fichier .{fmt} : {e}"
    except Exception as e:
        return f"⚠️ Erreur lors de la génération du fichier .{fmt} : {e}"


# Catalogue JSON (schéma pour le LLM
CATALOGUE_OUTILS = [
    {
        "type": "function",
        "function": {
            "name": "outil_meteo",
            "description": "Donne la météo complète et en temps réel d'une ville (température, humidité, vent, état du ciel).",
            "parameters": {
                "type": "object",
                "properties": {
                    "ville": {
                        "type": "string",
                        "description": "Nom de la ville (ex : Paris, London, Tokyo).",
                    }
                },
                "required": ["ville"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_wiki",
            "description": "Recherche et retourne un résumé Wikipédia sur n'importe quel sujet.",
            "parameters": {
                "type": "object",
                "properties": {
                    "sujet": {
                        "type": "string",
                        "description": "Sujet à rechercher sur Wikipédia (ex : Tour Eiffel, Albert Einstein).",
                    }
                },
                "required": ["sujet"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_argent",
            "description": "Convertit un montant d'une devise vers une autre avec les taux en temps réel.",
            "parameters": {
                "type": "object",
                "properties": {
                    "montant":      {"type": "number", "description": "Valeur à convertir (ex : 100)."},
                    "de_monnaie":   {"type": "string", "description": "Code ISO 4217 source (ex : EUR)."},
                    "vers_monnaie": {"type": "string", "description": "Code ISO 4217 cible  (ex : USD)."},
                },
                "required": ["montant", "de_monnaie", "vers_monnaie"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_duckduckgo",
            "description": (
                "Effectue une recherche web via DuckDuckGo et retourne les N premiers "
                "résultats (titre, URL, extrait). Utilise cet outil pour répondre à des "
                "questions nécessitant des informations récentes ou factuelles."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "La requête de recherche en langage naturel.",
                    },
                    "num_results": {
                        "type": "integer",
                        "description": "Nombre de résultats à retourner (défaut 5, max 10).",
                        "default": 5,
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_recup_page",
            "description": (
                "Télécharge et extrait le texte principal d'une page web à partir "
                "de son URL. Utile pour lire le contenu complet d'un article ou d'une page."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "L'URL complète de la page à récupérer.",
                    },
                    "max_chars": {
                        "type": "integer",
                        "description": "Nombre maximum de caractères à retourner (défaut 3000).",
                        "default": 3000,
                    },
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_datetime",
            "description": "Retourne la date et l'heure actuelles (horloge locale du serveur).",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "outil_generer_fichier",
            "description": (
                "Génère un fichier téléchargeable (txt, md, pdf, docx, xlsx) à partir d'un contenu textuel. "
                "Utilise cet outil quand l'utilisateur demande explicitement à exporter, sauvegarder, "
                "télécharger ou créer un fichier avec le résultat d'une réponse."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "contenu": {
                        "type": "string",
                        "description": "Contenu complet à mettre dans le fichier (texte brut ou markdown).",
                    },
                    "format": {
                        "type": "string",
                        "enum": ["txt", "md", "pdf", "docx", "xlsx"],
                        "description": "Format du fichier à générer.",
                    },
                    "nom_fichier": {
                        "type": "string",
                        "description": "Nom du fichier sans extension (ex : rapport_meteo, synthese_2024).",
                    },
                },
                "required": ["contenu", "format", "nom_fichier"],
            },
        },
    },
]

# Emojis d'affichage (utilisés dans l'interface web)
ICONES_OUTILS = {
    "outil_meteo":            "🌤️ Météo",
    "outil_wiki":             "📖 Wikipédia",
    "outil_argent":           "💱 Change",
    "outil_duckduckgo":       "🔍 Recherche web",
    "outil_recup_page":       "🌐 Lecture page",
    "outil_datetime":         "🕐 Date & Heure",
    "outil_generer_fichier":  "💾 Génération fichier",
}



# Fonctions utilitaires
# Retourne la liste des outils activés selon la section [tools] du .conf. Permet d'activer/désactiver chaque outil sans toucher au code.
def outils_actifs(conf: configparser.ConfigParser) -> list:
    mapping = {
        "enable_meteo":             "outil_meteo",
        "enable_wiki":              "outil_wiki",
        "enable_argent":            "outil_argent",
        "enable_duckduckgo":        "outil_duckduckgo",
        "enable_recup_page":        "outil_recup_page",
        "enable_datetime":          "outil_datetime",
        "enable_generer_fichier":   "outil_generer_fichier",
    }
    actifs = []
    for cle, nom in mapping.items():
        if conf.getboolean("tools", cle, fallback=True):
            actifs.extend([o for o in CATALOGUE_OUTILS if o["function"]["name"] == nom])
    return actifs

# Dispatcher central : appelle la bonne fonction selon le nom renvoyé par le LLM.
# Toujours utiliser cette fonction plutôt qu'appeler les outils directement.
def executer_outil(nom: str, args: dict) -> str:
    if nom == "outil_meteo":
        return outil_meteo(args["ville"])
    elif nom == "outil_wiki":
        return outil_wiki(args["sujet"])
    elif nom == "outil_argent":
        return outil_argent(args["montant"], args["de_monnaie"], args["vers_monnaie"])
    elif nom == "outil_duckduckgo":
        return outil_duckduckgo(args["query"], args.get("num_results", 5))
    elif nom == "outil_recup_page":
        return outil_recup_page(args["url"], args.get("max_chars", 3000))
    elif nom == "outil_datetime":
        return outil_datetime()
    elif nom == "outil_generer_fichier":
        return outil_generer_fichier(args["contenu"], args["format"], args["nom_fichier"])
    return f"⚠️ Outil inconnu : « {nom} »."
